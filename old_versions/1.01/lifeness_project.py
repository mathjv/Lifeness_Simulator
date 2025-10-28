# lifeness_Simulator

# Libraries
from __future__ import annotations
import os
import sys
import time
import math
import logging
from logging.handlers import RotatingFileHandler
from dataclasses import dataclass, field
from typing import List, Optional, Dict

# GUI, OpenGL
from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QPushButton, QListWidget,
    QTextEdit, QHBoxLayout, QVBoxLayout, QSplitter, QSlider, QMessageBox,
    QDialog, QFormLayout, QComboBox
)
from PySide6.QtGui import QFont, QAction
from PySide6.QtOpenGLWidgets import QOpenGLWidget
from OpenGL.GL import *
from OpenGL.GLU import *

# DOCX
from docx import Document
from docx.shared import Pt

# Paths base
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")
LOGS_DIR = os.path.join(BASE_DIR, "logs")
EXPORT_DIR = os.path.join(BASE_DIR, "export")
DOCX_SOURCE = os.path.join(BASE_DIR, "Doc1.docx")  # archivo subido por el usuario

# Asegurar carpetas
os.makedirs(ASSETS_DIR, exist_ok=True)
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
LOG_FILE = os.path.join(LOGS_DIR, "app.log")
logger = logging.getLogger("lifeness")
logger.setLevel(logging.DEBUG)
if not logger.handlers:
    fh = RotatingFileHandler(LOG_FILE, maxBytes=2_000_000, backupCount=3, encoding="utf-8")
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    ch = logging.StreamHandler()
    ch.setFormatter(fmt)
    logger.addHandler(ch)
logger.info("Inicio Lifeness Project - prototipo")

# ---------------------------------------------------------------------------
# Modelos de datos
# ---------------------------------------------------------------------------
@dataclass
class Enfermedad:
    id: str
    nombre: str
    categoria: str
    descripcion: str
    origen: Optional[str] = None
    transmision: Optional[str] = None
    prevencion: Optional[str] = None
    tratamiento: Optional[str] = None
    metadata: Dict = field(default_factory=dict)

@dataclass
class MetaProyecto:
    titulo: str = "Lifeness Simulator"
    autores: str = "Matthias Jimenez, Jose Herbas, Diego Guerron, Juandiego Vizuete"
    version: str = "1.01"
    fecha: str = ""
    idioma: str = "es"
    descripcion: str = "Simulador educativo."

class DocxParser:
    CATEGORIAS_BASE = ["Bacterias", "Virus", "Hongos", "Parásitos", "Priones"]

    def __init__(self, path: str = DOCX_SOURCE):
        self.path = path
        self.meta = MetaProyecto()
        self.enfermedades: List[Enfermedad] = []
        self.categorias = list(self.CATEGORIAS_BASE)

    def load(self) -> bool:
        logger.info("Intentando cargar documento fuente: %s", self.path)
        if not os.path.isfile(self.path):
            logger.warning("No se encontró Doc1.docx en la ruta del script. Se generarán casos demo.")
            self._generate_demo(min_cases=20)
            return False
        try:
            doc = Document(self.path)
        except Exception as e:
            logger.exception("Error abriendo Docx: %s", e)
            self._generate_demo(min_cases=20)
            return False

        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if paras:
            self.meta.titulo = paras[0][:120]
        for p in paras[:40]:
            low = p.lower()
            if "autor" in low or "autores" in low:
                self.meta.autores = p
            if "versi" in low or "version" in low:
                self.meta.version = p
            if "fecha" in low or "date" in low:
                self.meta.fecha = p
            if "idioma" in low:
                self.meta.idioma = "es" if "es" in low else "en"

        idx = 0
        for p in paras:
            text = p
            if ":" in text and len(text.split(":")[0]) < 50:
                name = text.split(":")[0].strip()
                descr = ":".join(text.split(":")[1:]).strip()
                cat = self._heuristic_categoria(text)
                eid = f"e{idx+1}"
                self.enfermedades.append(Enfermedad(id=eid, nombre=name, categoria=cat, descripcion=descr))
                idx += 1
            elif len(text) < 80 and any(c.isalpha() for c in text) and idx < 500:
                pos = paras.index(p)
                if pos + 1 < len(paras):
                    descr = paras[pos + 1][:300]
                    name = text
                    cat = self._heuristic_categoria(text)
                    eid = f"e{idx+1}"
                    self.enfermedades.append(Enfermedad(id=eid, nombre=name, categoria=cat, descripcion=descr))
                    idx += 1

        logger.info("Enfermedades extraídas: %d", len(self.enfermedades))
        # Si pocos casos, generar demo pero manteniendo los extraídos
        if len(self.enfermedades) < 20:
            self._generate_demo(min_cases=20)
        return True

    def _heuristic_categoria(self, text: str) -> str:
        """Asignar categoría basada en palabras clave del texto."""
        low = text.lower()
        if any(k in low for k in ["virus", "viral", "varicela", "hepatitis", "influenza", "dengue"]):
            return "Virus"
        if any(k in low for k in ["bacteria", "bacter", "tubercu", "pertuss"]):
            return "Bacterias"
        if any(k in low for k in ["hongo", "fung", "candida", "micosis"]):
            return "Hongos"
        if any(k in low for k in ["protozo", "parasi", "malaria"]):
            return "Parásitos"
        return "Virus"

    def _generate_demo(self, min_cases=20):
        """Genera ejemplos sintéticos claramente marcados como tal."""
        sample = [
            ("Influenza - ejemplo sintético", "Virus respiratorio estacional"),
            ("Dengue - ejemplo sintético", "Fiebre por arbovirus transmitido por mosquito"),
            ("Tuberculosis - ejemplo sintético", "Infección bacteriana pulmonar"),
            ("Otitis - ejemplo sintético", "Infección de oído medio"),
            ("Candidiasis - ejemplo sintético", "Infección fúngica superficial"),
        ]
        idx = len(self.enfermedades) + 1
        while len(self.enfermedades) < min_cases:
            s = sample[(idx - 1) % len(sample)]
            e = Enfermedad(
                id=f"demo_{idx}",
                nombre=f"{s[0]}",
                categoria=self._heuristic_categoria(s[0]),
                descripcion=f"{s[1]}. (Ejemplo sintético para demostración educativa).",
                origen="Ejemplo sintético",
                prevencion="Medidas generales: higiene, vacunación cuando aplique.",
                tratamiento="Protocolos generales (no dosis)."
            )
            self.enfermedades.append(e)
            idx += 1
        logger.info("Se completaron casos demo hasta: %d", len(self.enfermedades))

# ---------------------------------------------------------------------------
# OBJ Loader - crea display list (manejo robusto de errores)
# ---------------------------------------------------------------------------
class OBJ:
    """
    Cargador OBJ simple que crea un display list (GL) para render rápido.
    Maneja comentarios y faces trianguladas; si hay quads, intenta usarlas.
    """
    def __init__(self, filename: str):
        self.filename = filename
        self.vertices: List[List[float]] = []
        self.normals: List[List[float]] = []
        self.texcoords: List[List[float]] = []
        self.faces: List[List[tuple]] = []
        self.gl_list = None
        if os.path.isfile(filename):
            try:
                self._load_file(filename)
                logger.info("OBJ cargado: %s (v:%d f:%d)", filename, len(self.vertices), len(self.faces))
            except Exception as e:
                logger.exception("Error cargando OBJ %s: %s", filename, e)
        else:
            logger.warning("OBJ no encontrado: %s", filename)

    def _load_file(self, filename: str):
        """Parse básico del OBJ; tolerante a errores de encoding."""
        with open(filename, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                parts = line.strip().split()
                if not parts:
                    continue
                if parts[0] == 'v' and len(parts) >= 4:
                    self.vertices.append([float(parts[1]), float(parts[2]), float(parts[3])])
                elif parts[0] == 'vn' and len(parts) >= 4:
                    self.normals.append([float(parts[1]), float(parts[2]), float(parts[3])])
                elif parts[0] == 'vt' and len(parts) >= 3:
                    self.texcoords.append([float(parts[1]), float(parts[2])])
                elif parts[0] == 'f':
                    face = []
                    for v in parts[1:]:
                        vals = v.split('/')
                        vi = int(vals[0]) - 1 if vals[0] else -1
                        ti = int(vals[1]) - 1 if len(vals) > 1 and vals[1] else -1
                        ni = int(vals[2]) - 1 if len(vals) > 2 and vals[2] else -1
                        face.append((vi, ti, ni))
                    self.faces.append(face)

    def create_gl_list(self):
        """Crea la display list; evita recrearla si ya existe."""
        if self.gl_list is not None:
            return
        try:
            self.gl_list = glGenLists(1)
            glNewList(self.gl_list, GL_COMPILE)
            glEnable(GL_NORMALIZE)
            # Dibujar triángulos; si face tiene >3 vértices, usar GL_POLYGON como fallback
            for face in self.faces:
                if len(face) == 3:
                    glBegin(GL_TRIANGLES)
                    for (vi, ti, ni) in face:
                        if ni >= 0 and ni < len(self.normals):
                            glNormal3fv(self.normals[ni])
                        if ti >= 0 and ti < len(self.texcoords):
                            glTexCoord2fv(self.texcoords[ti])
                        if vi >= 0 and vi < len(self.vertices):
                            glVertex3fv(self.vertices[vi])
                    glEnd()
                else:
                    # fallback para quads o n-gons
                    glBegin(GL_POLYGON)
                    for (vi, ti, ni) in face:
                        if ni >= 0 and ni < len(self.normals):
                            glNormal3fv(self.normals[ni])
                        if ti >= 0 and ti < len(self.texcoords):
                            glTexCoord2fv(self.texcoords[ti])
                        if vi >= 0 and vi < len(self.vertices):
                            glVertex3fv(self.vertices[vi])
                    glEnd()
            glEndList()
        except Exception as e:
            logger.exception("Error creando GL list para %s: %s", self.filename, e)
            # limpiar lista si hay error
            try:
                if self.gl_list:
                    glDeleteLists(self.gl_list, 1)
                self.gl_list = None
            except Exception:
                pass

    def render(self):
        """Renderiza el modelo; si no existe GL list, la crea."""
        if self.gl_list is None:
            self.create_gl_list()
        if self.gl_list:
            try:
                glCallList(self.gl_list)
            except Exception as e:
                logger.exception("Error en glCallList: %s", e)

class GLHumanWidget(QOpenGLWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.model_male_path = os.path.join(ASSETS_DIR, "male.obj")
        self.model_female_path = os.path.join(ASSETS_DIR, "female.obj")
        self.model_male = OBJ(self.model_male_path) if os.path.isfile(self.model_male_path) else None
        self.model_female = OBJ(self.model_female_path) if os.path.isfile(self.model_female_path) else None
        self.current_model = self.model_male or self.model_female
        self.yaw = 0.0
        self.last_mouse_x = None
        self.zoom = -6.0
        self.bg_black = True
        self.reaction_id: Optional[str] = None
        self.reaction_start = 0.0
        self.timer = QtCore.QTimer(self)
        self.timer.timeout.connect(self.update)
        self.timer.start(30)

    # Interacción mouse
    def mousePressEvent(self, event):
        self.last_mouse_x = event.position().x()

    def mouseMoveEvent(self, event):
        if self.last_mouse_x is None:
            self.last_mouse_x = event.position().x()
            return
        dx = event.position().x() - self.last_mouse_x
        # sensibilidad moderada y suavizada
        self.yaw += dx * 0.3
        self.yaw = self.yaw % 360
        self.last_mouse_x = event.position().x()
        self.update()

    def mouseReleaseEvent(self, event):
        self.last_mouse_x = None

    def wheelEvent(self, event):
        delta = event.angleDelta().y() / 120.0
        self.zoom += delta * 0.6
        self.zoom = max(-20.0, min(-2.0, self.zoom))
        self.update()

    def mouseDoubleClickEvent(self, event):
        # toggle fondo negro/blanco y forzar repaint
        self.bg_black = not self.bg_black
        self.update()

    # OpenGL lifecycle
    def initializeGL(self):
        glEnable(GL_DEPTH_TEST)
        glEnable(GL_LIGHTING)
        glEnable(GL_LIGHT0)
        glLightfv(GL_LIGHT0, GL_POSITION, [4.0, 4.0, 10.0, 1.0])
        glEnable(GL_COLOR_MATERIAL)
        logger.debug("initializeGL OK")

    def resizeGL(self, w, h):
        glViewport(0, 0, w, h if h > 0 else 1)
        glMatrixMode(GL_PROJECTION)
        glLoadIdentity()
        gluPerspective(50.0, w / max(1.0, h), 0.1, 100.0)
        glMatrixMode(GL_MODELVIEW)

    def paintGL(self):
        # fondo
        if self.bg_black:
            glClearColor(0.05, 0.05, 0.06, 1.0)
        else:
            glClearColor(0.95, 0.95, 0.95, 1.0)
        glClear(GL_COLOR_BUFFER_BIT | GL_DEPTH_BUFFER_BIT)
        glLoadIdentity()
        glTranslatef(0.0, 0.0, self.zoom)
        glRotatef(self.yaw, 0.0, 1.0, 0.0)

        # efecto reacción pulsante
        alpha = 0.0
        if self.reaction_id:
            t = time.time() - self.reaction_start
            alpha = 0.25 + 0.5 * (0.5 + 0.5 * math.sin(t * 5.0))

        # coloreado global: si reacción -> tint roja
        if alpha > 0:
            glColor3f(1.0, 0.6 * (1 - alpha), 0.6 * (1 - alpha))
        else:
            glColor3f(0.9, 0.88, 0.85)

        # render modelo con fallback seguro
        try:
            if self.current_model:
                self.current_model.render()
            else:
                # fallback: dibujar un torso simple con primitivas
                self._draw_placeholder_human()
        except Exception as e:
            logger.exception("Error al renderizar modelo GL: %s", e)
            self._draw_placeholder_human()

    def _draw_placeholder_human(self):
        # dibujo minimalista para evitar crash si OBJ falla
        glPushMatrix()
        glTranslatef(0.0, 0.6, 0.0)
        quad = gluNewQuadric()
        gluSphere(quad, 0.25, 16, 12)
        gluDeleteQuadric(quad)
        glPopMatrix()
        # torso
        glPushMatrix()
        glTranslatef(0.0, -0.25, 0.0)
        glScalef(1.0, 1.6, 0.5)
        self._draw_cube(0.4)
        glPopMatrix()

    def _draw_cube(self, s):
        glBegin(GL_QUADS)
        # front
        glVertex3f(-s, -s, s)
        glVertex3f(s, -s, s)
        glVertex3f(s, s, s)
        glVertex3f(-s, s, s)
        # back
        glVertex3f(-s, -s, -s)
        glVertex3f(-s, s, -s)
        glVertex3f(s, s, -s)
        glVertex3f(s, -s, -s)
        glEnd()

    def apply_reaction(self, enfermedad_id: Optional[str]):
        """Inicia efecto visual asociado a una enfermedad (no clínico)."""
        self.reaction_id = enfermedad_id
        if enfermedad_id:
            self.reaction_start = time.time()
        self.update()

    def set_gender_model(self, gender: str):
        """Permitir cambiar entre modelos male/female si existen en assets."""
        if gender.lower().startswith("m") and self.model_male:
            self.current_model = self.model_male
        elif gender.lower().startswith("f") and self.model_female:
            self.current_model = self.model_female
        self.update()

# ---------------------------------------------------------------------------
# Settings Dialog
# ---------------------------------------------------------------------------
class SettingsDialog(QDialog):
    """Dialogo de ajustes: brillo, contraste, fullscreen, idioma."""
    def __init__(self, parent=None, current_lang="es"):
        super().__init__(parent)
        self.setWindowTitle("Ajustes")
        self.resize(420, 200)
        layout = QFormLayout(self)
        self.sldr_brightness = QSlider(QtCore.Qt.Horizontal)
        self.sldr_brightness.setRange(0, 100); self.sldr_brightness.setValue(50)
        self.sldr_contrast = QSlider(QtCore.Qt.Horizontal)
        self.sldr_contrast.setRange(0, 100); self.sldr_contrast.setValue(50)
        self.chk_fullscreen = QtWidgets.QCheckBox("Pantalla completa")
        self.cmb_lang = QComboBox(); self.cmb_lang.addItems(["es", "en"]); self.cmb_lang.setCurrentText(current_lang)
        layout.addRow("Brillo", self.sldr_brightness)
        layout.addRow("Contraste", self.sldr_contrast)
        layout.addRow(self.chk_fullscreen)
        layout.addRow("Idioma", self.cmb_lang)
        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addRow(btns)

    def values(self):
        return {
            "brightness": self.sldr_brightness.value(),
            "contrast": self.sldr_contrast.value(),
            "fullscreen": self.chk_fullscreen.isChecked(),
            "lang": self.cmb_lang.currentText()
        }

# ---------------------------------------------------------------------------
# Report Generator (.docx) con estructura APA7 (básica)
# ---------------------------------------------------------------------------
class ReportGenerator:
    """Genera Reporte_Proyecto_Lifeness.docx con índice por categoría y enfermedades."""
    def __init__(self, meta: MetaProyecto, enfermedades: List[Enfermedad], out_path: str = os.path.join(EXPORT_DIR, "Reporte_Proyecto_Lifeness.docx")):
        self.meta = meta
        self.enfermedades = enfermedades
        self.out_path = out_path

    def generate(self) -> str:
        """Genera el documento y retorna la ruta."""
        logger.info("Generando reporte docx en: %s", self.out_path)
        doc = Document()
        # Portada
        h = doc.add_heading(self.meta.titulo, level=1)
        h.alignment = 1
        p = doc.add_paragraph(f"Autores: {self.meta.autores}\n")
        p.add_run(f"Institución: {self.meta.titulo}\n")
        p.add_run(f"Versión: {self.meta.version}    Fecha: {self.meta.fecha}\n")
        doc.add_page_break()
        # Resumen ejecutivo
        doc.add_heading("Resumen ejecutivo", level=2)
        doc.add_paragraph(self.meta.descripcion or "Resumen no disponible.")
        # Advertencia
        doc.add_heading("Advertencia", level=2)
        doc.add_paragraph("Material educativo. No sustituye la evaluación médica profesional.")
        # Mini-libro por categoría
        categories: Dict[str, List[Enfermedad]] = {}
        for e in self.enfermedades:
            categories.setdefault(e.categoria, []).append(e)
        for cat, items in categories.items():
            doc.add_heading(f"Capítulo: {cat}", level=2)
            for ent in items:
                doc.add_heading(ent.nombre, level=3)
                doc.add_paragraph(ent.descripcion or "Descripción no disponible.")
                doc.add_paragraph(f"Prevención: {ent.prevencion or 'No disponible.'}")
                doc.add_paragraph(f"Tratamiento (general): {ent.tratamiento or 'Protocolos generales (sin dosis).'}")
        # Anexos
        doc.add_heading("Anexos: Casos y tablas", level=2)
        for i, e in enumerate(self.enfermedades[:200], 1):
            doc.add_paragraph(f"{i}. {e.nombre} — {e.categoria}")
        # Bibliografía (placeholder APA7)
        doc.add_heading("Bibliografía (APA7 - muestra)", level=2)
        doc.add_paragraph("World Health Organization. (2025). Título. https://www.who.int/")
        doc.save(self.out_path)
        logger.info("Reporte guardado: %s", self.out_path)
        return self.out_path

# ---------------------------------------------------------------------------
# Main Window: composición de UI + lógica (comentarios funcionales)
# ---------------------------------------------------------------------------
class MainWindow(QMainWindow):
    """Ventana principal que compone la UI solicitada y orquesta interacciones."""
    def __init__(self, parser: DocxParser):
        super().__init__()
        self.parser = parser
        self.meta = parser.meta
        self.setWindowTitle(f"{self.meta.titulo} — Lifeness Project")
        self.resize(1400, 820)
        # Layout: splitter horizontal (izq categorías, der centro+right)
        main_split = QSplitter(QtCore.Qt.Horizontal)
        # Panel izquierdo: lista vertical con botones (imitando estilo Office Save-as)
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(6, 6, 6, 6)
        # Botones grandes (5) - estilo uniforme
        self.btn_info = QPushButton("Info / Actualizar")
        self.btn_sim = QPushButton("Simulador y Enfermedades")
        self.btn_ajustes = QPushButton("Ajustes")
        self.btn_report = QPushButton("Generar Reporte")
        self.btn_exit = QPushButton("Salir")
        for btn in (self.btn_info, self.btn_sim, self.btn_ajustes, self.btn_report, self.btn_exit):
            btn.setMinimumHeight(56)
            btn.setFont(QFont("Arial", 11))
            btn.setToolTip("Haga clic para ejecutar la acción")
            left_layout.addWidget(btn)
        left_layout.addStretch()
        # Scroll area for more buttons (if needed)
        left_scroll = QtWidgets.QScrollArea()
        left_scroll.setWidget(left_widget)
        left_scroll.setWidgetResizable(True)
        left_scroll.setFixedWidth(320)
        main_split.addWidget(left_scroll)

        # Centro: GL widget y timeline debajo
        center_widget = QWidget()
        center_layout = QVBoxLayout(center_widget)
        center_layout.setContentsMargins(6, 6, 6, 6)
        # Título arriba
        title_lbl = QLabel(self.meta.titulo)
        title_lbl.setFont(QFont("Arial Black", 18))
        title_lbl.setAlignment(QtCore.Qt.AlignCenter)
        center_layout.addWidget(title_lbl)
        # GL viewer
        self.gl_widget = GLHumanWidget()
        center_layout.addWidget(self.gl_widget, 1)
        # timeline controls
        timeline_bar = QWidget()
        t_layout = QHBoxLayout(timeline_bar)
        self.btn_play = QPushButton("Play")
        self.btn_pause = QPushButton("Pause")
        self.cmb_speed = QComboBox(); self.cmb_speed.addItems(["0.5x", "1x", "2x"]); self.cmb_speed.setCurrentText("1x")
        self.tslider = QSlider(QtCore.Qt.Horizontal); self.tslider.setRange(0, 100); self.tslider.setValue(0)
        t_layout.addWidget(self.btn_play); t_layout.addWidget(self.btn_pause); t_layout.addWidget(self.cmb_speed); t_layout.addWidget(self.tslider)
        center_layout.addWidget(timeline_bar)
        main_split.addWidget(center_widget)

        # Right: OMS info + enfermedades list panel (scrollable)
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(6, 6, 6, 6)
        self.txt_oms = QTextEdit()
        self.txt_oms.setReadOnly(True)
        # cargar un texto informativo (incluye advertencia)
        self.txt_oms.setHtml("<b>OMS - Recursos y referencias</b><p>Este simulador es material educativo. No sustituye atención médica profesional.</p>")
        right_layout.addWidget(self.txt_oms, 1)
        # listado de enfermedades por categoria (se actualizará)
        self.lst_enf = QListWidget()
        right_layout.addWidget(QLabel("Enfermedades (seleccione):"))
        right_layout.addWidget(self.lst_enf, 1)
        self.btn_tratamiento = QPushButton("Tratamiento")
        right_layout.addWidget(self.btn_tratamiento)
        main_split.addWidget(right_widget)
        main_split.setStretchFactor(1, 2)

        # Central widget
        central = QWidget()
        central_layout = QHBoxLayout(central)
        central_layout.addWidget(main_split)
        self.setCentralWidget(central)

        # Conexiones
        self.btn_info.clicked.connect(self.show_info)
        self.btn_sim.clicked.connect(self.show_sim_categories)
        self.btn_ajustes.clicked.connect(self.show_settings)
        self.btn_exit.clicked.connect(self.confirm_exit)
        self.btn_report.clicked.connect(self.generate_report)
        self.lst_enf.itemClicked.connect(self.on_enfermedad_selected)
        self.btn_tratamiento.clicked.connect(self.on_tratamiento_clicked)
        self.btn_play.clicked.connect(self.on_play)
        self.btn_pause.clicked.connect(self.on_pause)
        self.cmb_speed.currentTextChanged.connect(self.on_speed_change)
        self.tslider.valueChanged.connect(self.on_timeline_move)

        # Shortcuts
        QtGui_short_f11 = QAction(self)
        QtGui_short_f11.setShortcut("F11")
        QtGui_short_f11.triggered.connect(self.toggle_fullscreen)
        self.addAction(QtGui_short_f11)
        QtGui_short_esc = QAction(self)
        QtGui_short_esc.setShortcut("Esc")
        QtGui_short_esc.triggered.connect(self.on_escape)
        self.addAction(QtGui_short_esc)

        # Poblamos categorias y enfermedades iniciales desde parser
        self.populate_categories()
        # Estado timeline
        self.timeline_playing = False
        self.timeline_speed = 1.0
        self.timeline_timer = QtCore.QTimer()
        self.timeline_timer.timeout.connect(self.advance_timeline)

    # ---- UI helpers y acciones ----
    def populate_categories(self):
        """Poblar la lista de enfermedades derecha con las primeras 10 (o por categoria)."""
        # por defecto muestro las primeras 10 de parser
        self.lst_enf.clear()
        for e in self.parser.enfermedades[:50]:
            self.lst_enf.addItem(f"{e.nombre} [{e.categoria}]")
            item = self.lst_enf.item(self.lst_enf.count() - 1)
            item.setToolTip(e.descripcion[:200])

    def show_info(self):
        """Muestra modal con metadatos y resumen del proyecto."""
        dlg = QMessageBox(self)
        dlg.setWindowTitle("Información - Lifeness Project")
        info = (
            f"<b>{self.meta.titulo}</b><br>"
            f"Autores: {self.meta.autores}<br>"
            f"Versión: {self.meta.version}    Fecha: {self.meta.fecha}<br>"
            f"Idioma: {self.meta.idioma}<br><br>"
            f"{self.meta.descripcion}<br><br>"
            "Nota: Material educativo. No sustituye evaluación médica profesional."
        )
        dlg.setText(info)
        dlg.exec()

    def show_sim_categories(self):
        """Limpia pantalla central y muestra un menú de categorías (UI simplificada)."""
        # Mostrar categorías en lista central (simple) — aquí reusamos lst_enf para selección
        self.lst_enf.clear()
        for cat in self.parser.categorias:
            self.lst_enf.addItem(f"[Categoria] {cat}")
            item = self.lst_enf.item(self.lst_enf.count() - 1)
            item.setToolTip(f"Mostrar enfermedades de {cat}")
        # Se limpia GL view (opcional)
        self.gl_widget.apply_reaction(None)

    def show_settings(self):
        dlg = SettingsDialog(self, current_lang=self.meta.idioma)
        if dlg.exec() == QDialog.Accepted:
            vals = dlg.values()
            logger.info("Ajustes aplicados: %s", vals)
            # aplicar fullscreen
            if vals["fullscreen"]:
                self.showFullScreen()
            else:
                self.showNormal()
            # idioma switch minimal
            self.meta.idioma = vals["lang"]

    def confirm_exit(self):
        resp = QMessageBox.question(self, "Confirmar salida", "¿Desea salir?", QMessageBox.Yes | QMessageBox.No)
        if resp == QMessageBox.Yes:
            QMessageBox.information(self, "Gracias", "Gracias por usar Lifeness Project.\nMaterial educativo.")
            logger.info("Usuario cerró la aplicación.")
            QApplication.instance().quit()

    def generate_report(self):
        rg = ReportGenerator(self.meta, self.parser.enfermedades)
        path = rg.generate()
        QMessageBox.information(self, "Reporte generado", f"Reporte guardado en: {path}")

    def on_enfermedad_selected(self, item):
        text = item.text()
        # Buscar enfermedad por nombre (heurística)
        name = text.split(" [")[0]
        e = next((x for x in self.parser.enfermedades if x.nombre.startswith(name)), None)
        if not e:
            # si item es categoria
            if text.startswith("[Categoria]"):
                cat = text.replace("[Categoria] ", "")
                # filtrar enfermedades por categoria
                filt = [x for x in self.parser.enfermedades if x.categoria.lower() == cat.lower()]
                self.lst_enf.clear()
                for ent in filt[:10]:
                    self.lst_enf.addItem(f"{ent.nombre} [{ent.categoria}]")
                return
            return
        # mostrar ficha debajo del centro (usamos txt_oms para visibilidad)
        ficha = (
            f"<b>{e.nombre}</b><br>"
            f"Categoría: {e.categoria}<br>"
            f"Descripción: {e.descripcion}<br>"
            f"Prevención: {e.prevencion or 'No disponible.'}<br>"
            f"<i>Fuente: {os.path.basename(self.parser.path) if os.path.isfile(self.parser.path) else 'Demo'}</i><br>"
            f"<p style='color:darkred;'>Advertencia: Material educativo - No sustituye la evaluación médica profesional.</p>"
        )
        self.txt_oms.setHtml(ficha)
        # aplicar reacción visual en el modelo (no clínico)
        self.gl_widget.apply_reaction(e.id)

    def on_tratamiento_clicked(self):
        # Cambiar vista para mostrar tratamientos (simplificado)
        sel = self.lst_enf.currentItem()
        if not sel:
            QMessageBox.information(self, "Tratamiento", "Seleccione una enfermedad en la lista.")
            return
        name = sel.text().split(" [")[0]
        e = next((x for x in self.parser.enfermedades if x.nombre.startswith(name)), None)
        if not e:
            QMessageBox.information(self, "Tratamiento", "Enfermedad no encontrada.")
            return
        texto = (
            f"<b>Tratamientos sugeridos (generales)</b><br>"
            f"{e.nombre}<br>"
            f"Protocolos generales: {e.tratamiento or 'Manejo clínico estándar (sin dosis).'}<br>"
            "<i>Nota: no se incluyen dosificaciones. Consulte guías oficiales y personal médico autorizado.</i>"
        )
        self.txt_oms.setHtml(texto)

    # Timeline controls
    def on_play(self):
        if not self.timeline_playing:
            self.timeline_timer.start(int(1000 / self.timeline_speed))
            self.timeline_playing = True

    def on_pause(self):
        if self.timeline_playing:
            self.timeline_timer.stop()
            self.timeline_playing = False

    def on_speed_change(self, txt):
        self.timeline_speed = 0.5 if txt.startswith("0.5") else (2.0 if txt.startswith("2") else 1.0)
        if self.timeline_playing:
            self.timeline_timer.start(int(1000 / self.timeline_speed))

    def advance_timeline(self):
        val = self.tslider.value()
        if val < self.tslider.maximum():
            self.tslider.setValue(val + 1)
        else:
            self.tslider.setValue(0)
            # loop or stop
            # self.on_pause()

    def on_timeline_move(self, val):
        # Si el timeline avanza, se podrían cambiar fases de reacción; ejemplo: a val>30 cambia intensidad
        if val < 33:
            self.gl_widget.apply_reaction(None)
        elif val < 66:
            self.gl_widget.apply_reaction("phase_mid")
        else:
            self.gl_widget.apply_reaction("phase_high")

    # Other UI actions
    def toggle_fullscreen(self):
        if self.isFullScreen():
            self.showNormal()
        else:
            self.showFullScreen()

    def on_escape(self):
        if self.isFullScreen():
            self.showNormal()

# ---------------------------------------------------------------------------
# Splash screen con transiciones precisas
# ---------------------------------------------------------------------------
class SplashWindow(QWidget):
    """Splash con fade-in 0.6s, display 5s total, fade-out 0.6s y blackout 900µs antes de mostrar main."""
    def __init__(self, meta: MetaProyecto):
        super().__init__()
        self.meta = meta
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.resize(800, 420)
        layout = QVBoxLayout(self)
        title = QLabel(self.meta.titulo)
        title.setFont(QFont("Arial Black", 22))
        title.setStyleSheet("color: white;")
        title.setAlignment(QtCore.Qt.AlignCenter)
        sub = QLabel(f"{self.meta.descripcion}\nAutores: {self.meta.autores}\nVersión: {self.meta.version} {self.meta.fecha}")
        sub.setStyleSheet("color: #ddd;")
        sub.setAlignment(QtCore.Qt.AlignCenter)
        layout.addStretch()
        layout.addWidget(title)
        layout.addWidget(sub)
        layout.addStretch()
        self.opacity = QtWidgets.QGraphicsOpacityEffect(self)
        self.setGraphicsEffect(self.opacity)

    def play(self):
        """Efecto fade-in -> hold -> fade-out -> blackout 900 microseg -> done."""
        # Fade-in 0.6s
        steps = 30
        for i in range(steps):
            self.opacity.setOpacity(i / (steps - 1))
            QApplication.processEvents()
            time.sleep(0.6 / steps)
        # Hold until total 5s (we already used 0.6)
        hold = 5.0 - 0.6 - 0.6
        time.sleep(max(0.0, hold))
        # Fade-out 0.6s
        for i in range(steps):
            self.opacity.setOpacity(1.0 - (i / (steps - 1)))
            QApplication.processEvents()
            time.sleep(0.6 / steps)
        # Blackout 900 microsegundos (~0.0009 s) - practically instantaneous; simulate tiny pause
        time.sleep(0.001)

# ---------------------------------------------------------------------------
# Utilities: generar README y build script
# ---------------------------------------------------------------------------
# def ensure_readme_and_build():
#     """Crea README.md y build_exe.bat con comandos PyInstaller si no existen."""
#     readme = os.path.join(BASE_DIR, "README.md")
#     if not os.path.isfile(readme):
#         with open(readme, "w", encoding="utf-8") as f:
#             f.write("# Lifeness Project - Prototipo\n\n")
#             f.write("Requisitos:\n- Python 3.11+\n- pip install PySide6 PyOpenGL python-docx Pillow numpy\n\n")
#             f.write("Ejecución:\n```\npython lifeness_project.py\n```\n\n")
#             f.write("Empaquetado (Windows):\n```\npy -3.11 -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt\npyinstaller --onefile --windowed --hidden-import=OpenGL.GL --hidden-import=OpenGL.GLU --add-data \"assets;assets\" lifeness_project.py\n```\n")
#     bat = os.path.join(BASE_DIR, "build_exe.bat")
#     if not os.path.isfile(bat):
#         with open(bat, "w", encoding="utf-8") as f:
#             f.write("py -3.11 -m venv venv\r\nvenv\\Scripts\\activate\r\npip install PySide6 PyOpenGL python-docx Pillow numpy\r\npyinstaller --onefile --windowed --hidden-import=OpenGL.GL --hidden-import=OpenGL.GLU --add-data \"assets;assets\" lifeness_project.py\r\n")
#     reqs = os.path.join(BASE_DIR, "requirements.txt")
#     if not os.path.isfile(reqs):
#         with open(reqs, "w", encoding="utf-8") as f:
#             f.write("PySide6\nPyOpenGL\npython-docx\nPillow\nnumpy\n")
#     logger.info("README y scripts de build asegurados")
# ---------------------------------------------------------------------------
# Tests unitarios básicos
# ---------------------------------------------------------------------------
def run_unit_tests() -> bool:
    """Ejecuta pruebas unitarias simples para parser y generador de reportes."""
    import unittest

    class ParserTest(unittest.TestCase):
        def test_parser_handles_missing_file(self):
            p = DocxParser(path=os.path.join(BASE_DIR, "nonexistent.docx"))
            ok = p.load()
            self.assertFalse(ok)
            self.assertGreaterEqual(len(p.enfermedades), 20)

    class ReportTest(unittest.TestCase):
        def test_report_generation(self):
            meta = MetaProyecto(titulo="Test", autores="A", version="0.1", fecha="2025")
            e = [Enfermedad(id="t1", nombre="Test disease", categoria="Virus", descripcion="Desc")]
            rg = ReportGenerator(meta, e, out_path=os.path.join(EXPORT_DIR, "tmp_report.docx"))
            path = rg.generate()
            self.assertTrue(os.path.isfile(path))
            os.remove(path)

    suite = unittest.TestLoader().loadTestsFromTestCase(ParserTest)
    suite2 = unittest.TestLoader().loadTestsFromTestCase(ReportTest)
    alltests = unittest.TestSuite([suite, suite2])
    result = unittest.TextTestRunner(verbosity=2).run(alltests)
    return result.wasSuccessful()

# ---------------------------------------------------------------------------
# Main: orquesta splash -> main window
# ---------------------------------------------------------------------------
def main():
    #ensure_readme_and_build()
    parser = DocxParser()
    parser.load()  # usa Doc1.docx si está presente; si no, genera demo
    app = QApplication(sys.argv)
    # crear y mostrar splash con transiciones (bloqueante corto)
    splash = SplashWindow(parser.meta)
    splash.show()
    # play splash in blocking way (small sleeps)
    splash.play()
    splash.close()
    # main window
    window = MainWindow(parser)
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
